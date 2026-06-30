import os

from pypdf import PdfReader
from docx import Document as DocxDocument

from langchain_core.documents import Document

from services.splitter import split_text


def extract_text(file_path: str) -> str:
    """
    Extract text based on file type.
    """

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".txt":
        return read_txt(file_path)

    elif ext == ".pdf":
        return read_pdf(file_path)

    elif ext == ".docx":
        return read_docx(file_path)

    else:
        raise ValueError(f"Unsupported document type: {ext}")


def read_txt(file_path: str) -> str:

    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def read_pdf(file_path: str) -> str:

    reader = PdfReader(file_path)

    pages = []

    for page in reader.pages:

        text = page.extract_text()

        if text:
            pages.append(text)

    return "\n".join(pages)


def read_docx(file_path: str) -> str:

    document = DocxDocument(file_path)

    paragraphs = [
        para.text
        for para in document.paragraphs
        if para.text.strip()
    ]

    return "\n".join(paragraphs)


def create_documents(file_path: str):
    """
    Convert document into LangChain Documents.
    """

    text = extract_text(file_path)

    chunks = split_text(text)

    documents = []

    filename = os.path.basename(file_path)

    for idx, chunk in enumerate(chunks):

        documents.append(

            Document(

                page_content=chunk,

                metadata={
                    "source": filename,
                    "chunk": idx
                }

            )

        )

    return documents